"""
Word文档处理工具，支持将Word文档转换为Markdown格式
"""
import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import tempfile
import zipfile
import xml.etree.ElementTree as ET
import re

logger = logging.getLogger(__name__)

class WordProcessor:
    """Word文档处理器，用于将Word文档转换为Markdown"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
        print(f"🔧 Word处理器初始化完成")
        print(f"   支持格式: {', '.join(self.supported_formats)}")
    
    def is_word_document(self, file_path: str) -> bool:
        """
        检查文件是否为Word文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否为Word文档
        """
        return Path(file_path).suffix.lower() in self.supported_formats
    
    def docx_to_markdown_native(self, docx_path: str, output_dir: str) -> str:
        """
        使用原生方法将DOCX转换为Markdown（解析XML结构）
        
        Args:
            docx_path: DOCX文件路径
            output_dir: 输出目录
            
        Returns:
            输出文件路径
        """
        docx_path = Path(docx_path)
        output_dir = Path(output_dir)
        
        print(f"\n📄 使用原生方法处理DOCX文档")
        print("=" * 80)
        print(f"   文件名: {docx_path.name}")
        print(f"   文件大小: {docx_path.stat().st_size / 1024:.1f} KB")
        print(f"   输出目录: {output_dir}")
        
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX文件不存在: {docx_path}")
        
        try:
            # 创建输出目录
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 解压DOCX文件
            print(f"\n📦 步骤1: 解压DOCX文件")
            print("-" * 40)
            
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                
                with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_path)
                
                # 读取document.xml
                document_xml_path = temp_path / "word" / "document.xml"
                if not document_xml_path.exists():
                    raise ValueError("无效的DOCX文件：缺少document.xml")
                
                print(f"✅ DOCX文件解压成功")
                
                # 解析XML内容
                print(f"\n🔍 步骤2: 解析XML内容")
                print("-" * 40)
                
                markdown_content = self._parse_docx_xml(document_xml_path)
                
                print(f"✅ XML解析完成")
                print(f"   提取内容长度: {len(markdown_content)} 字符")
                
                # 保存Markdown文件
                print(f"\n💾 步骤3: 保存Markdown文件")
                print("-" * 40)
                
                output_path = output_dir / docx_path.with_suffix(".md").name
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                
                print(f"✅ 文件保存成功")
                print(f"   输出路径: {output_path}")
                print(f"   内容长度: {len(markdown_content)} 字符")
                
                # 计算段落数（避免f-string中的反斜杠）
                paragraph_count = markdown_content.count('\n\n')
                print(f"   估算段落数: {paragraph_count} 个")
                
                print("=" * 80)
                print(f"🎉 原生方法处理完成!")
                print(f"   输出文件: {output_path}")
                print(f"   处理方式: 原生XML解析")
                
                return str(output_path)
                
        except Exception as e:
            print(f"\n❌ 原生方法处理失败: {e}")
            logger.error(f"Native DOCX processing failed {docx_path.name}: {e}")
            raise
    
    def _parse_docx_xml(self, document_xml_path: Path) -> str:
        """
        解析DOCX的document.xml文件，提取文本内容并转换为Markdown
        
        Args:
            document_xml_path: document.xml文件路径
            
        Returns:
            Markdown格式的内容
        """
        try:
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            # 定义命名空间
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            markdown_lines = []
            
            # 查找所有段落
            paragraphs = root.findall('.//w:p', namespaces)
            
            for para in paragraphs:
                paragraph_text = self._extract_paragraph_text(para, namespaces)
                
                if paragraph_text.strip():
                    # 检查是否为标题
                    if self._is_heading(para, namespaces):
                        heading_level = self._get_heading_level(para, namespaces)
                        markdown_lines.append(f"{'#' * heading_level} {paragraph_text}")
                    else:
                        markdown_lines.append(paragraph_text)
                    
                    markdown_lines.append("")  # 添加空行
            
            # 处理表格
            tables = root.findall('.//w:tbl', namespaces)
            for table in tables:
                table_markdown = self._extract_table_markdown(table, namespaces)
                if table_markdown:
                    markdown_lines.extend(table_markdown)
                    markdown_lines.append("")
            
            return "\n".join(markdown_lines).strip()
            
        except Exception as e:
            logger.error(f"Error parsing DOCX XML: {e}")
            raise
    
    def _extract_paragraph_text(self, paragraph, namespaces) -> str:
        """提取段落文本"""
        text_parts = []
        
        # 查找所有文本运行(run)
        runs = paragraph.findall('.//w:r', namespaces)
        
        for run in runs:
            # 检查格式
            is_bold = run.find('.//w:b', namespaces) is not None
            is_italic = run.find('.//w:i', namespaces) is not None
            
            # 提取文本
            text_elements = run.findall('.//w:t', namespaces)
            for text_elem in text_elements:
                text = text_elem.text or ""
                
                # 应用格式
                if is_bold and is_italic:
                    text = f"***{text}***"
                elif is_bold:
                    text = f"**{text}**"
                elif is_italic:
                    text = f"*{text}*"
                
                text_parts.append(text)
        
        return "".join(text_parts)
    
    def _is_heading(self, paragraph, namespaces) -> bool:
        """检查段落是否为标题"""
        # 检查样式
        style_elem = paragraph.find('.//w:pStyle', namespaces)
        if style_elem is not None:
            style_val = style_elem.get(f'{{{namespaces["w"]}}}val', '')
            return 'heading' in style_val.lower() or 'title' in style_val.lower()
        
        return False
    
    def _get_heading_level(self, paragraph, namespaces) -> int:
        """获取标题级别"""
        style_elem = paragraph.find('.//w:pStyle', namespaces)
        if style_elem is not None:
            style_val = style_elem.get(f'{{{namespaces["w"]}}}val', '')
            
            # 尝试从样式名中提取级别
            import re
            match = re.search(r'(\d+)', style_val)
            if match:
                level = int(match.group(1))
                return min(level, 6)  # 最多6级标题
        
        return 1  # 默认一级标题
    
    def _extract_table_markdown(self, table, namespaces) -> list:
        """提取表格并转换为Markdown格式"""
        markdown_lines = []
        rows = table.findall('.//w:tr', namespaces)
        
        if not rows:
            return markdown_lines
        
        table_data = []
        
        for row in rows:
            cells = row.findall('.//w:tc', namespaces)
            row_data = []
            
            for cell in cells:
                cell_text = ""
                cell_paragraphs = cell.findall('.//w:p', namespaces)
                
                for para in cell_paragraphs:
                    para_text = self._extract_paragraph_text(para, namespaces)
                    if para_text.strip():
                        cell_text += para_text + " "
                
                row_data.append(cell_text.strip())
            
            if row_data:
                table_data.append(row_data)
        
        if table_data:
            # 生成Markdown表格
            if len(table_data) > 0:
                # 表头
                header = "| " + " | ".join(table_data[0]) + " |"
                markdown_lines.append(header)
                
                # 分隔线
                separator = "| " + " | ".join(["---"] * len(table_data[0])) + " |"
                markdown_lines.append(separator)
                
                # 数据行
                for row in table_data[1:]:
                    if len(row) == len(table_data[0]):  # 确保列数一致
                        data_row = "| " + " | ".join(row) + " |"
                        markdown_lines.append(data_row)
        
        return markdown_lines
    
    def docx_to_markdown_python_docx(self, docx_path: str, output_dir: str) -> str:
        """
        使用python-docx库将DOCX转换为Markdown
        
        Args:
            docx_path: DOCX文件路径
            output_dir: 输出目录
            
        Returns:
            输出文件路径
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
        
        docx_path = Path(docx_path)
        output_dir = Path(output_dir)
        
        print(f"\n📚 使用python-docx库处理DOCX文档")
        print("=" * 80)
        print(f"   文件名: {docx_path.name}")
        
        try:
            # 创建输出目录
            output_dir.mkdir(parents=True, exist_ok=True)
            
            print(f"📖 正在解析DOCX文件...")
            doc = Document(docx_path)
            
            print(f"🔄 转换为Markdown格式...")
            markdown_lines = []
            
            # 处理段落
            for para in doc.paragraphs:
                if para.text.strip():
                    # 检查样式
                    style_name = para.style.name.lower()
                    
                    if 'heading' in style_name:
                        # 提取标题级别
                        level = 1
                        if 'heading 1' in style_name:
                            level = 1
                        elif 'heading 2' in style_name:
                            level = 2
                        elif 'heading 3' in style_name:
                            level = 3
                        elif 'heading 4' in style_name:
                            level = 4
                        elif 'heading 5' in style_name:
                            level = 5
                        elif 'heading 6' in style_name:
                            level = 6
                        
                        markdown_lines.append(f"{'#' * level} {para.text}")
                    else:
                        # 处理格式化文本
                        formatted_text = self._format_paragraph_runs(para)
                        markdown_lines.append(formatted_text)
                    
                    markdown_lines.append("")  # 添加空行
            
            # 处理表格
            for table in doc.tables:
                table_markdown = self._convert_table_to_markdown(table)
                markdown_lines.extend(table_markdown)
                markdown_lines.append("")
            
            markdown_content = "\n".join(markdown_lines).strip()
            
            print(f"💾 保存Markdown文件...")
            output_path = output_dir / docx_path.with_suffix(".md").name
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print("=" * 80)
            print(f"🎉 python-docx处理完成!")
            print(f"   输出文件: {output_path}")
            print(f"   内容长度: {len(markdown_content)} 字符")
            print(f"   处理方式: python-docx库")
            
            return str(output_path)
            
        except Exception as e:
            print(f"❌ python-docx处理失败: {e}")
            logger.error(f"python-docx processing failed {docx_path.name}: {e}")
            raise
    
    def _format_paragraph_runs(self, paragraph):
        """格式化段落中的文本运行"""
        formatted_text = ""
        
        for run in paragraph.runs:
            text = run.text
            
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            formatted_text += text
        
        return formatted_text
    
    def _convert_table_to_markdown(self, table):
        """将表格转换为Markdown格式"""
        markdown_lines = []
        
        if not table.rows:
            return markdown_lines
        
        # 获取表格数据
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if table_data:
            # 表头
            header = "| " + " | ".join(table_data[0]) + " |"
            markdown_lines.append(header)
            
            # 分隔线
            separator = "| " + " | ".join(["---"] * len(table_data[0])) + " |"
            markdown_lines.append(separator)
            
            # 数据行
            for row in table_data[1:]:
                if len(row) == len(table_data[0]):
                    data_row = "| " + " | ".join(row) + " |"
                    markdown_lines.append(data_row)
        
        return markdown_lines
    
    def doc_to_markdown_antiword(self, doc_path: str, output_dir: str) -> str:
        """
        使用antiword工具将DOC转换为文本，然后转换为Markdown
        
        Args:
            doc_path: DOC文件路径
            output_dir: 输出目录
            
        Returns:
            输出文件路径
        """
        import subprocess
        
        doc_path = Path(doc_path)
        output_dir = Path(output_dir)
        
        print(f"\n🔧 使用antiword工具处理DOC文档")
        print("=" * 80)
        print(f"   文件名: {doc_path.name}")
        
        try:
            # 检查antiword是否可用
            result = subprocess.run(['antiword', '-v'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode != 0:
                raise RuntimeError("antiword工具不可用，请安装: sudo apt-get install antiword")
            
            # 创建输出目录
            output_dir.mkdir(parents=True, exist_ok=True)
            
            print(f"📖 正在使用antiword提取文本...")
            
            # 使用antiword提取文本
            result = subprocess.run(['antiword', str(doc_path)], 
                                  capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0:
                raise RuntimeError(f"antiword处理失败: {result.stderr}")
            
            text_content = result.stdout
            
            print(f"🔄 转换为Markdown格式...")
            
            # 简单的文本到Markdown转换
            markdown_content = self._text_to_markdown(text_content)
            
            print(f"💾 保存Markdown文件...")
            output_path = output_dir / doc_path.with_suffix(".md").name
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print("=" * 80)
            print(f"🎉 antiword处理完成!")
            print(f"   输出文件: {output_path}")
            print(f"   内容长度: {len(markdown_content)} 字符")
            print(f"   处理方式: antiword工具")
            
            return str(output_path)
            
        except subprocess.TimeoutExpired:
            print(f"❌ antiword处理超时")
            raise
        except Exception as e:
            print(f"❌ antiword处理失败: {e}")
            logger.error(f"antiword processing failed {doc_path.name}: {e}")
            raise
    
    def _text_to_markdown(self, text_content: str) -> str:
        """
        将纯文本转换为基本的Markdown格式
        
        Args:
            text_content: 纯文本内容
            
        Returns:
            Markdown格式的内容
        """
        lines = text_content.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                markdown_lines.append("")
                continue
            
            # 简单的标题检测（全大写或以数字开头）
            if (line.isupper() and len(line) < 100) or \
               (re.match(r'^\d+\.?\s+[A-Z]', line) and len(line) < 100):
                markdown_lines.append(f"## {line}")
            else:
                markdown_lines.append(line)
            
            markdown_lines.append("")
        
        return "\n".join(markdown_lines).strip()
    
    def word_to_markdown(self, word_path: str, output_dir: str, method: str = "auto") -> str:
        """
        将Word文档转换为Markdown
        
        Args:
            word_path: Word文档路径
            output_dir: 输出目录
            method: 转换方法 ("auto", "native", "python-docx", "antiword")
            
        Returns:
            输出文件路径
        """
        word_path = Path(word_path)
        
        print(f"\n📄 Word文档处理任务开始")
        print(f"   文件: {word_path.name}")
        print(f"   方法: {method}")
        
        if not word_path.exists():
            raise FileNotFoundError(f"Word文档不存在: {word_path}")
        
        if not self.is_word_document(str(word_path)):
            raise ValueError(f"不支持的文件格式: {word_path.suffix}")
        
        # 根据文件类型和方法选择处理方式
        if method == "auto":
            if word_path.suffix.lower() == '.docx':
                # 优先使用python-docx，失败时使用原生方法
                try:
                    return self.docx_to_markdown_python_docx(str(word_path), output_dir)
                except ImportError:
                    print(f"⚠️  python-docx不可用，切换到原生方法...")
                    return self.docx_to_markdown_native(str(word_path), output_dir)
                except Exception as e:
                    print(f"⚠️  python-docx处理失败，切换到原生方法: {e}")
                    return self.docx_to_markdown_native(str(word_path), output_dir)
            else:  # .doc文件
                return self.doc_to_markdown_antiword(str(word_path), output_dir)
        
        elif method == "native":
            if word_path.suffix.lower() == '.docx':
                return self.docx_to_markdown_native(str(word_path), output_dir)
            else:
                raise ValueError("原生方法只支持DOCX文件")
        
        elif method == "python-docx":
            if word_path.suffix.lower() == '.docx':
                return self.docx_to_markdown_python_docx(str(word_path), output_dir)
            else:
                raise ValueError("python-docx方法只支持DOCX文件")
        
        elif method == "antiword":
            if word_path.suffix.lower() == '.doc':
                return self.doc_to_markdown_antiword(str(word_path), output_dir)
            else:
                raise ValueError("antiword方法只支持DOC文件")
        
        else:
            raise ValueError(f"不支持的转换方法: {method}")